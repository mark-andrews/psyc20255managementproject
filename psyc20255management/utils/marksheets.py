"""Utilities for creating and processing marksheets, etc.

"""
#=============================================================================
# Standard library imports
#=============================================================================
import os
import re
import zipfile

#=============================================================================
# Third party imports
#=============================================================================
from docx import Document
from bs4 import BeautifulSoup

#=============================================================================
# Imports of homespun packages
#=============================================================================
from ernst.emisc import assertTrue, assertEqual

#=============================================================================
# Local imports
#=============================================================================
from .. import conf

#================================ End Imports ================================

def process_completed_marksheets(completed_marksheets_dirname,
                                 sequence_name='Experimental',
                                 marksheet_fname_pattern=None):


    '''Return a list of processed and checked marksheets that were found in the
    completed_marksheets_dirname. 

    '''

    completed_marksheets = []
    for completed_marksheet\
            in list_completed_marksheets(completed_marksheets_dirname,
                                         marksheet_fname_pattern):
        
        try:
            (student_name, 
             student_id, 
             marker_name, 
             marker_email, 
             grade) = MarksheetModel.get_marksheet_vital_details(
                     marksheet_filename=completed_marksheet['filepath'],
                     sequence_name = sequence_name)
            
            assertEqual(student_name, completed_marksheet['student_name'])
            assertEqual(student_id, completed_marksheet['student_id'])
            assertTrue(grade in conf.grades)
            
            completed_marksheets.append(
                [student_name, 
                 student_id, 
                 marker_name,
                 marker_email,
                 grade, 
                 completed_marksheet['filepath']]
            )
        except AssertionError:
            print('Bad trouble with %s.' % completed_marksheet)
            raise

    return completed_marksheets


def list_completed_marksheets(completed_marking_dirname,
                              marksheet_fname_pattern=None):
    
    '''Return a list of (hopefully completed) marksheets from a
    `completed_marking` directory. Marksheets are defined as files that match a
    particular regular expression, which is recorded in conf.
    
    The list is a list of small dictionaries, with keys 
    * student_name, which is taken from the filename 
    * student_id, which is also taken from the filename
    * filename, the filename
    * filepath, the full and absolute path to the file
    '''

    # We need this here because we have changed the default
    # marksheet_fname_pattern from using commas to using double underscores.
    # If we want to process old marksheets or new marksheets, we'll need to
    # change this pattern.
    if marksheet_fname_pattern is None:
        marksheet_fname_pattern = conf.marksheet_fname_pattern

    def abspath(dirname, fname):
        'Convenience function just to save some typing'
        return os.path.abspath(os.path.join(completed_marking_dirname, fname))
    
    completed_marksheets = []
    
    for fname in os.listdir(completed_marking_dirname):
        
        pattern_match = marksheet_fname_pattern.match(fname)
        
        if pattern_match: 

            student_name, student_id = pattern_match.groups()
            
            completed_marksheets.append(
                    dict(student_name = student_name,
                         student_id = student_id,
                         filename = fname,
                         filepath = abspath(completed_marking_dirname, fname)
                         )
                    )
        else:
            print('Did not match: %s' % fname)
    
    return completed_marksheets


def get_grade_from_marksheet(marksheet):
    
    ''' 
    Extract the assigned grade from the marksheet. Return as a string that
    should be from the `grades` list.

    The grade assigned to a report is chosen from a Word dropdown list.  The
    chosen value in a Word dropdown list is not obtainable using python-docx,
    and so we must pull out the xml inside the zipfile (a docx file is a zip
    file) and parse that using beautiful soup. The relevant tag for the
    dropdown list is `sdt` (I think), and we're assuming there'll be only one
    such tag in the file. The chosen value in the dropdown list is the value of
    the `sdtContent` tag (again, I think).

    '''
    
    document = zipfile.ZipFile(marksheet)
    xml_data = document.read('word/document.xml')
    document.close()
    soup = BeautifulSoup(xml_data, 'xml')
    grade = soup.find('sdt').find('sdtContent').text

    assert grade in conf.grades,\
            'Grade %s not in grades list %s.' % (grade, ' ,'.join(conf.grades))

    return grade


class MarksheetModel(object):

    '''A model of a a marksheet document

    This is used from reading and parsing marksheets, and creating new
    marksheets from a generic version.

    The top of the document, i.e. the first 6 paragraphs, have a particular
    structure:

    * The first line in the document should be a heading labelled 'PSYC20255:
    (Something) Sequence', e.g. 'PSYC20255: Psychometrics Sequence' 
    * Then there should be blank.
    * Then "Student name: (Jane Doe)"
    * Then "Student ID: (N0666007)"
    * Then "Marker name: (John Foo)"
    * Then "Marker email: (foo.bar@ntu.ac.uk)"

    where the info in the parentheses are variable.

    '''

    title_par_index = 0
    null_par_index = 1
    student_name_par_index = 2
    student_ID_par_index = 3
    marker_name_par_index = 4
    marker_email_par_index = 5
    marker_grade_par_index = 6
    
    student_name_label = 'Student name'
    student_ID_label = 'Student ID'
    marker_name_label = 'Marker name'
    marker_email_label = 'Marker email'
    marker_grade_label = 'Grade'
    
    student_name_pattern = re.compile(r'%s: (.*)$' % student_name_label)
    student_id_pattern = re.compile(r'%s: (.*)$' % student_ID_label)
    marker_name_pattern = re.compile(r'%s: (.*)$' % marker_name_label)
    marker_email_pattern = re.compile(r'%s: (.*)$' % marker_email_label)
    marker_grade_pattern = re.compile(r'%s: (.*)$' % marker_grade_label)

    def __init__(self, document_name, sequence_name='Experimental'):
        
        self.document_name = document_name
        self.document = Document(self.document_name)
        self.sequence_name = 'PSYC20255: %s Sequence' % sequence_name
        self.validate()
 
    #### Class methods ####
    @classmethod
    def get_marksheet_vital_details(cls, marksheet_filename, sequence_name):
        '''Return the "vital details" of a marksheet.

        This function uses the `extract_vital_details` instance method of the
        `MarksheetModel` class.
        '''

        marksheet_model = cls(marksheet_filename, sequence_name)
        return marksheet_model.extract_vital_details()

    ###########
       
    def get_paragraph_contents(self):
        'Return the text contents of each paragraph of the document.'
        return [p.text for p in self.document.paragraphs]
        
    def validate(self):
        ''' This validates the starting paragraph structure. Anything
        unexpected will raise an Assertion error.'''
        
        P = self.get_paragraph_contents()
        
        assertEqual(P[self.title_par_index], self.sequence_name)
        assertEqual(P[self.null_par_index], '')
        assertTrue(self.student_name_pattern.match(P[self.student_name_par_index]))
        assertTrue(self.student_id_pattern.match(P[self.student_ID_par_index]))
        assertTrue(self.marker_name_pattern.match(P[self.marker_name_par_index]))
        assertTrue(self.marker_email_pattern.match(P[self.marker_email_par_index]))
        assertTrue(self.marker_grade_pattern.match(P[self.marker_grade_par_index]))
        
    def extract_vital_details(self):

        '''Return the following as a tuple:

        * Student name
        * Student ID
        * Marker name
        * Marker email
        * Assigned grade (i.e. value of dropdown)

        '''
        
        P = self.get_paragraph_contents()

        # This BS is to deal with what happens if someone messes with the
        # dropdown menu. 
        try:
            grade = get_grade_from_marksheet(self.document_name)
        except AttributeError:
            print("Probably can't read dropdown in %s." % self.document_name)
            grade = self.marker_grade_pattern.match(P[self.marker_grade_par_index]).groups()[0]
                    
            try:
                assert grade in conf.grades
            except:
                grade = grade.strip().replace('.','').upper()

            assert grade in conf.grades,\
                'Grade %s not in grades list %s.' % (grade, ' ,'.join(conf.grades))

        
        return (self.student_name_pattern.match(P[self.student_name_par_index]).groups()[0],
                self.student_id_pattern.match(P[self.student_ID_par_index]).groups()[0],
                self.marker_name_pattern.match(P[self.marker_name_par_index]).groups()[0],
                self.marker_email_pattern.match(P[self.marker_email_par_index]).groups()[0],
                grade)
    
    def make_new_marksheet(self, 
                          student_name,
                          student_ID, 
                          marker_name, 
                          marker_email,
                          new_marksheet_name):

        '''
        Create a new marksheet named `new_marksheet_name` where 
        * the student name is `student_name`
        * student ID is `student_ID`
        * marker name is `marker_name`, 
        * marker email is `marker_email`.

        '''
        
        paragraphs = self.document.paragraphs
        
        paragraphs[self.student_name_par_index].text\
                = '%s: %s' % (self.student_name_label, student_name)
        paragraphs[self.student_ID_par_index].text\
                = '%s: %s' % (self.student_ID_label, student_ID)
        paragraphs[self.marker_name_par_index].text\
                = '%s: %s' % (self.marker_name_label, marker_name)
        paragraphs[self.marker_email_par_index].text\
                = '%s: %s' % (self.marker_email_label, marker_email)
        
        self.document.save(new_marksheet_name)


